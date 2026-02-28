from __future__ import annotations

from datetime import datetime, timezone
import hashlib
import json
import re
import requests
from uuid import uuid4

from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from pydantic import BaseModel, Field
from sqlalchemy import select
from sqlalchemy.orm import Session

from ainern2d_shared.ainer_db_models.content_models import Chapter, Novel
from ainern2d_shared.ainer_db_models.enum_models import KBBindType, RagScope, RagSourceType
from ainern2d_shared.ainer_db_models.governance_models import CreativePolicyStack
from ainern2d_shared.ainer_db_models.governance_models import PersonaPack, PersonaPackVersion
from ainern2d_shared.ainer_db_models.preview_models import PersonaDatasetBinding, PersonaIndexBinding
from ainern2d_shared.ainer_db_models.provider_models import ModelProvider
from ainern2d_shared.ainer_db_models.rag_models import KbVersion, RagCollection, RagDocument

from app.api.deps import get_db
from app.services.telegram_notify import notify_telegram_event

router = APIRouter(prefix="/api/v1/rag", tags=["rag"])


class RagCollectionCreateRequest(BaseModel):
    tenant_id: str
    project_id: str
    name: str
    language_code: str = "zh"
    description: str | None = None
    novel_id: str | None = None
    tags_json: list[str] = Field(default_factory=list)
    bind_type: str | None = None  # role / persona / novel / global
    bind_id: str | None = None    # 对应的 role_id / persona_pack_id / novel_id


class RagCollectionResponse(BaseModel):
    id: str
    tenant_id: str
    project_id: str
    name: str
    language_code: str | None = None
    description: str | None = None
    tags_json: list[str] = Field(default_factory=list)
    bind_type: str | None = None
    bind_id: str | None = None


class KbVersionCreateRequest(BaseModel):
    tenant_id: str
    project_id: str
    version_name: str
    status: str = "draft"
    recipe_id: str | None = None
    embedding_model_profile_id: str | None = None
    release_note: str | None = None


class KbVersionResponse(BaseModel):
    id: str
    collection_id: str
    version_name: str
    status: str
    recipe_id: str | None = None
    embedding_model_profile_id: str | None = None
    release_note: str | None = None


class PersonaPackCreateRequest(BaseModel):
    tenant_id: str
    project_id: str
    name: str
    role_id: str | None = None  # 关联的职业 ID，用于继承 Role KB
    description: str | None = None
    tags_json: list[str] = Field(default_factory=list)


class PersonaPackResponse(BaseModel):
    id: str
    tenant_id: str
    project_id: str
    name: str
    role_id: str | None = None
    description: str | None = None
    tags_json: list[str] = Field(default_factory=list)


class PersonaVersionCreateRequest(BaseModel):
    tenant_id: str
    project_id: str
    version_name: str
    style_json: dict = Field(default_factory=dict)
    voice_json: dict = Field(default_factory=dict)
    camera_json: dict = Field(default_factory=dict)


class PersonaVersionResponse(BaseModel):
    id: str
    persona_pack_id: str
    version_name: str
    style_json: dict = Field(default_factory=dict)
    voice_json: dict = Field(default_factory=dict)
    camera_json: dict = Field(default_factory=dict)


class PersonaBindingRequest(BaseModel):
    tenant_id: str
    project_id: str
    dataset_collection_ids: list[str] = Field(default_factory=list)
    kb_version_ids: list[str] = Field(default_factory=list)
    binding_role: str = "primary"


class PersonaBindingResponse(BaseModel):
    persona_pack_version_id: str
    dataset_bindings: list[str] = Field(default_factory=list)
    index_bindings: list[str] = Field(default_factory=list)


class PersonaPreviewRequest(BaseModel):
    tenant_id: str
    project_id: str
    persona_pack_version_id: str
    query: str
    top_k: int = Field(default=5, ge=1, le=20)


class PersonaPreviewChunk(BaseModel):
    doc_id: str
    collection_id: str | None = None
    kb_version_id: str | None = None
    title: str | None = None
    source_type: str
    source_id: str | None = None
    score: float
    snippet: str


class PersonaPreviewResponse(BaseModel):
    persona_pack_version_id: str
    query: str
    top_k: int
    chunks: list[PersonaPreviewChunk] = Field(default_factory=list)


class KnowledgePackBootstrapRequest(BaseModel):
    tenant_id: str
    project_id: str
    role_id: str
    pack_name: str | None = None
    template_key: str | None = None
    language_code: str = "zh"
    default_knowledge_scope: str = "style_rule"


class KnowledgePackBootstrapResponse(BaseModel):
    pack_id: str
    tenant_id: str
    project_id: str
    role_id: str
    pack_name: str
    collection_id: str
    kb_version_id: str
    scope: str
    created_documents: int
    chunk_count: int
    extracted_terms: list[str] = Field(default_factory=list)
    updated_at: str | None = None


class KnowledgePackItemResponse(BaseModel):
    pack_id: str
    tenant_id: str
    project_id: str
    role_id: str
    pack_name: str
    collection_id: str
    kb_version_id: str
    scope: str
    status: str
    created_documents: int
    extracted_terms: list[str] = Field(default_factory=list)
    updated_at: str | None = None


class KnowledgeImportRequest(BaseModel):
    tenant_id: str
    project_id: str
    collection_id: str
    kb_version_id: str | None = None
    source_format: str = "txt"
    source_name: str
    content_text: str
    role_ids: list[str] = Field(default_factory=list)
    language_code: str = "zh"
    scope: str = "chapter"


class KnowledgeImportJobResponse(BaseModel):
    import_job_id: str
    tenant_id: str
    project_id: str
    collection_id: str
    kb_version_id: str | None = None
    source_name: str
    source_format: str
    status: str
    created_documents: int
    deduplicated_documents: int
    chunk_count: int
    extracted_terms: list[str] = Field(default_factory=list)
    affected_roles: list[str] = Field(default_factory=list)
    knowledge_change_report: dict = Field(default_factory=dict)
    updated_at: str | None = None


class BinaryImportJobResponse(BaseModel):
    """TASK_CARD_26: 二进制文件导入任务"""
    import_job_id: str
    tenant_id: str
    project_id: str
    collection_id: str
    file_name: str
    file_format: str  # pdf / xlsx / txt / docx
    file_size_bytes: int
    status: str  # uploading / processing / completed / failed
    progress_percent: int = 0
    extracted_text_preview: str = ""
    extracted_pages: int = 0
    extracted_tables: int = 0
    extracted_images: int = 0
    created_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    updated_at: datetime = Field(default_factory=lambda: datetime.now(timezone.utc))
    error_message: str | None = None


def _knowledge_pack_stack_name(pack_id: str) -> str:
    return f"knowledge_pack:{pack_id}"


def _import_job_stack_name(import_job_id: str) -> str:
    return f"rag_import_job:{import_job_id}"


def _upsert_stack(
    db: Session,
    *,
    tenant_id: str,
    project_id: str,
    stack_name: str,
    payload: dict,
    trace_prefix: str,
) -> None:
    row = db.execute(
        select(CreativePolicyStack).where(
            CreativePolicyStack.tenant_id == tenant_id,
            CreativePolicyStack.project_id == project_id,
            CreativePolicyStack.name == stack_name,
            CreativePolicyStack.deleted_at.is_(None),
        )
    ).scalars().first()
    if row is None:
        row = CreativePolicyStack(
            id=f"policy_{uuid4().hex}",
            tenant_id=tenant_id,
            project_id=project_id,
            trace_id=f"tr_{trace_prefix}_{uuid4().hex[:12]}",
            correlation_id=f"cr_{trace_prefix}_{uuid4().hex[:12]}",
            idempotency_key=f"idem_{trace_prefix}_{uuid4().hex[:8]}",
            name=stack_name,
            status="active",
            stack_json=payload,
        )
        db.add(row)
    else:
        row.status = "active"
        row.stack_json = payload


def _chunk_text(content: str, chunk_size: int = 480) -> list[str]:
    paragraphs = [item.strip() for item in re.split(r"\n{2,}", content) if item.strip()]
    chunks: list[str] = []
    for paragraph in paragraphs:
        if len(paragraph) <= chunk_size:
            chunks.append(paragraph)
            continue
        start = 0
        while start < len(paragraph):
            end = min(len(paragraph), start + chunk_size)
            chunks.append(paragraph[start:end].strip())
            start = end
    return [item for item in chunks if item]


def _extract_terms(content: str, limit: int = 16) -> list[str]:
    latin_terms = [item.lower() for item in re.findall(r"[A-Za-z][A-Za-z0-9_\-]{2,30}", content)]
    cjk_terms = re.findall(r"[\u4e00-\u9fff]{2,6}", content)
    scores: dict[str, int] = {}
    for term in latin_terms + cjk_terms:
        scores[term] = scores.get(term, 0) + 1
    sorted_items = sorted(scores.items(), key=lambda item: (-item[1], item[0]))
    return [item[0] for item in sorted_items[:limit]]


def _knowledge_template_sections(role_id: str, template_key: str | None) -> list[dict]:
    role = (role_id or "").strip().lower()
    template = (template_key or "").strip().lower()
    if template:
        role = template

    default_sections = [
        {
            "title": "通用制作流程",
            "content": "输出必须先给结论再给依据，结构化列出风险、假设和下一步，默认返回可执行清单。",
            "source_type": RagSourceType.policy,
        },
        {
            "title": "质量检查清单",
            "content": "检查角色一致性、时序连续性、资产命名规范、文化约束冲突、提示词可执行性。",
            "source_type": RagSourceType.rule,
        },
    ]
    sections_by_role: dict[str, list[dict]] = {
        "director": [
            {
                "title": "导演镜头语法",
                "content": "镜头设计遵循建立镜头-关系镜头-动作镜头三段式，优先保持人物方位、轴线与情绪曲线一致。",
                "source_type": RagSourceType.rule,
            },
            {
                "title": "导演输出格式",
                "content": "输出分镜时必须包含 shot_id、时长、机位、主体、动作、转场与风险注记，支持 patch 重生成。",
                "source_type": RagSourceType.policy,
            },
        ],
        "script_supervisor": [
            {
                "title": "场记连续性规则",
                "content": "记录每镜服化道状态、台词改动和出入场时间，任何变化必须写入 continuity notes。",
                "source_type": RagSourceType.rule,
            },
        ],
        "art": [
            {
                "title": "美术设定基线",
                "content": "场景道具需与世界观时间线一致，避免时代错位元素，输出包含材质、色板与风格参考。",
                "source_type": RagSourceType.rule,
            },
        ],
        "lighting": [
            {
                "title": "灯光计划模板",
                "content": "每镜给出主光/辅光/轮廓光方案，明确色温区间、布光方向与对比度目标。",
                "source_type": RagSourceType.rule,
            },
        ],
        "stunt": [
            {
                "title": "动作分解SOP",
                "content": "按预备动作-冲突动作-收势动作分解，并标注安全边界、替身需求和慢动作节点。",
                "source_type": RagSourceType.rule,
            },
        ],
        "translator": [
            {
                "title": "翻译术语基准",
                "content": "先维护术语一致性，再处理语气与语域；输出中保留人名实体映射与文化注释。",
                "source_type": RagSourceType.rule,
            },
        ],
    }
    return sections_by_role.get(role, []) + default_sections


def _scope_from_text(scope: str) -> RagScope:
    try:
        return RagScope(scope)
    except ValueError:
        return RagScope.chapter


def _source_type_from_format(source_format: str) -> RagSourceType:
    mapping = {
        "txt": RagSourceType.note,
        "text": RagSourceType.note,
        "md": RagSourceType.note,
        "markdown": RagSourceType.note,
        "pdf": RagSourceType.policy,
        "excel": RagSourceType.rule,
        "xlsx": RagSourceType.rule,
        "csv": RagSourceType.rule,
        "tsv": RagSourceType.rule,
    }
    return mapping.get(source_format.strip().lower(), RagSourceType.note)


def _normalize_import_content(content_text: str, source_format: str) -> str:
    fmt = source_format.strip().lower()
    if fmt in {"excel", "xlsx", "csv", "tsv"}:
        rows = [item.strip() for item in content_text.splitlines() if item.strip()]
        rendered_rows: list[str] = []
        for row in rows:
            if "\t" in row:
                cells = [cell.strip() for cell in row.split("\t")]
            elif "," in row:
                cells = [cell.strip() for cell in row.split(",")]
            else:
                cells = [row]
            rendered_rows.append(" | ".join([cell for cell in cells if cell]))
        return "\n".join(rendered_rows)
    return content_text


def _build_knowledge_pack_item(
    payload: dict,
    *,
    pack_id: str,
    tenant_id: str,
    project_id: str,
) -> KnowledgePackItemResponse:
    return KnowledgePackItemResponse(
        pack_id=pack_id,
        tenant_id=tenant_id,
        project_id=project_id,
        role_id=str(payload.get("role_id") or ""),
        pack_name=str(payload.get("pack_name") or ""),
        collection_id=str(payload.get("collection_id") or ""),
        kb_version_id=str(payload.get("kb_version_id") or ""),
        scope=str(payload.get("scope") or "style_rule"),
        status=str(payload.get("status") or "ready"),
        created_documents=int(payload.get("created_documents") or 0),
        extracted_terms=list(payload.get("extracted_terms") or []),
        updated_at=payload.get("updated_at"),
    )


def _build_import_job_item(
    payload: dict,
    *,
    import_job_id: str,
    tenant_id: str,
    project_id: str,
) -> KnowledgeImportJobResponse:
    return KnowledgeImportJobResponse(
        import_job_id=import_job_id,
        tenant_id=tenant_id,
        project_id=project_id,
        collection_id=str(payload.get("collection_id") or ""),
        kb_version_id=payload.get("kb_version_id"),
        source_name=str(payload.get("source_name") or ""),
        source_format=str(payload.get("source_format") or "txt"),
        status=str(payload.get("status") or "completed"),
        created_documents=int(payload.get("created_documents") or 0),
        deduplicated_documents=int(payload.get("deduplicated_documents") or 0),
        chunk_count=int(payload.get("chunk_count") or 0),
        extracted_terms=list(payload.get("extracted_terms") or []),
        affected_roles=list(payload.get("affected_roles") or []),
        knowledge_change_report=dict(payload.get("knowledge_change_report") or {}),
        updated_at=payload.get("updated_at"),
    )


@router.post("/collections", response_model=RagCollectionResponse, status_code=201)
def create_collection(body: RagCollectionCreateRequest, db: Session = Depends(get_db)) -> RagCollectionResponse:
    # Parse bind_type
    parsed_bind_type: KBBindType | None = None
    if body.bind_type:
        bt = body.bind_type.strip().lower()
        if bt == "global":
            bt = "global_"
        try:
            parsed_bind_type = KBBindType(bt.rstrip("_")) if bt != "global_" else KBBindType.global_
        except ValueError:
            raise HTTPException(status_code=400, detail=f"invalid bind_type: {body.bind_type}")

    row = RagCollection(
        id=f"rag_collection_{uuid4().hex}",
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        trace_id=f"tr_rag_collection_{uuid4().hex[:12]}",
        correlation_id=f"cr_rag_collection_{uuid4().hex[:12]}",
        idempotency_key=f"idem_rag_collection_{body.name}_{uuid4().hex[:8]}",
        novel_id=body.novel_id,
        name=body.name,
        version="v1",
        language_code=body.language_code,
        description=body.description,
        tags_json=body.tags_json,
        bind_type=parsed_bind_type,
        bind_id=body.bind_id,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return RagCollectionResponse(
        id=row.id,
        tenant_id=row.tenant_id,
        project_id=row.project_id,
        name=row.name,
        language_code=row.language_code,
        description=row.description,
        tags_json=row.tags_json or [],
        bind_type=row.bind_type.value if row.bind_type else None,
        bind_id=row.bind_id,
    )


@router.get("/collections", response_model=list[RagCollectionResponse])
def list_collections(
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    keyword: str | None = Query(default=None),
    language_code: str | None = Query(default=None),
    bind_type: str | None = Query(default=None),
    bind_id: str | None = Query(default=None),
    db: Session = Depends(get_db),
) -> list[RagCollectionResponse]:
    stmt = select(RagCollection).where(
        RagCollection.tenant_id == tenant_id,
        RagCollection.project_id == project_id,
        RagCollection.deleted_at.is_(None),
    )
    if language_code:
        stmt = stmt.where(RagCollection.language_code == language_code)
    if keyword:
        like_value = f"%{keyword.strip()}%"
        stmt = stmt.where(RagCollection.name.ilike(like_value))
    if bind_type:
        bt = bind_type.strip().lower()
        if bt == "global":
            stmt = stmt.where(RagCollection.bind_type == KBBindType.global_)
        else:
            try:
                stmt = stmt.where(RagCollection.bind_type == KBBindType(bt))
            except ValueError:
                pass
    if bind_id:
        stmt = stmt.where(RagCollection.bind_id == bind_id)
    rows = db.execute(stmt.order_by(RagCollection.created_at.desc())).scalars().all()
    return [
        RagCollectionResponse(
            id=row.id,
            tenant_id=row.tenant_id,
            project_id=row.project_id,
            name=row.name,
            language_code=row.language_code,
            description=row.description,
            tags_json=row.tags_json or [],
            bind_type=row.bind_type.value if row.bind_type else None,
            bind_id=row.bind_id,
        )
        for row in rows
    ]


@router.post("/collections/{collection_id}/kb-versions", response_model=KbVersionResponse, status_code=201)
def create_kb_version(
    collection_id: str,
    body: KbVersionCreateRequest,
    db: Session = Depends(get_db),
) -> KbVersionResponse:
    collection = db.get(RagCollection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="collection not found")

    row = KbVersion(
        id=f"kb_version_{uuid4().hex}",
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        trace_id=f"tr_kb_{uuid4().hex[:12]}",
        correlation_id=f"cr_kb_{uuid4().hex[:12]}",
        idempotency_key=f"idem_kb_{collection_id}_{body.version_name}_{uuid4().hex[:8]}",
        collection_id=collection_id,
        version_name=body.version_name,
        status=body.status,
        recipe_id=body.recipe_id,
        embedding_model_profile_id=body.embedding_model_profile_id,
        release_note=body.release_note,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return KbVersionResponse(
        id=row.id,
        collection_id=row.collection_id,
        version_name=row.version_name,
        status=row.status,
        recipe_id=row.recipe_id,
        embedding_model_profile_id=row.embedding_model_profile_id,
        release_note=row.release_note,
    )


@router.get("/collections/{collection_id}/kb-versions", response_model=list[KbVersionResponse])
def list_kb_versions(
    collection_id: str,
    status: str | None = Query(default=None),
    db: Session = Depends(get_db),
) -> list[KbVersionResponse]:
    stmt = select(KbVersion).where(
        KbVersion.collection_id == collection_id,
        KbVersion.deleted_at.is_(None),
    )
    if status:
        stmt = stmt.where(KbVersion.status == status)
    rows = db.execute(stmt.order_by(KbVersion.created_at.desc())).scalars().all()
    return [
        KbVersionResponse(
            id=row.id,
            collection_id=row.collection_id,
            version_name=row.version_name,
            status=row.status,
            recipe_id=row.recipe_id,
            embedding_model_profile_id=row.embedding_model_profile_id,
            release_note=row.release_note,
        )
        for row in rows
    ]


@router.post("/persona-packs", response_model=PersonaPackResponse, status_code=201)
def create_persona_pack(body: PersonaPackCreateRequest, db: Session = Depends(get_db)) -> PersonaPackResponse:
    row = PersonaPack(
        id=f"persona_pack_{uuid4().hex}",
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        trace_id=f"tr_persona_pack_{uuid4().hex[:12]}",
        correlation_id=f"cr_persona_pack_{uuid4().hex[:12]}",
        idempotency_key=f"idem_persona_pack_{body.name}_{uuid4().hex[:8]}",
        name=body.name,
        role_id=body.role_id,
        description=body.description,
        tags_json=body.tags_json,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return PersonaPackResponse(
        id=row.id,
        tenant_id=row.tenant_id,
        project_id=row.project_id,
        name=row.name,
        role_id=row.role_id,
        description=row.description,
        tags_json=row.tags_json or [],
    )


@router.get("/persona-packs", response_model=list[PersonaPackResponse])
def list_persona_packs(
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    keyword: str | None = Query(default=None),
    db: Session = Depends(get_db),
) -> list[PersonaPackResponse]:
    stmt = select(PersonaPack).where(
        PersonaPack.tenant_id == tenant_id,
        PersonaPack.project_id == project_id,
        PersonaPack.deleted_at.is_(None),
    )
    if keyword:
        like_value = f"%{keyword.strip()}%"
        stmt = stmt.where(PersonaPack.name.ilike(like_value))
    rows = db.execute(stmt.order_by(PersonaPack.created_at.desc())).scalars().all()
    return [
        PersonaPackResponse(
            id=row.id,
            tenant_id=row.tenant_id,
            project_id=row.project_id,
            name=row.name,
            role_id=row.role_id,
            description=row.description,
            tags_json=row.tags_json or [],
        )
        for row in rows
    ]


@router.post("/persona-packs/{persona_pack_id}/versions", response_model=PersonaVersionResponse, status_code=201)
def create_persona_version(
    persona_pack_id: str,
    body: PersonaVersionCreateRequest,
    db: Session = Depends(get_db),
) -> PersonaVersionResponse:
    pack = db.get(PersonaPack, persona_pack_id)
    if pack is None:
        raise HTTPException(status_code=404, detail="persona pack not found")

    row = PersonaPackVersion(
        id=f"persona_ver_{uuid4().hex}",
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        trace_id=f"tr_persona_ver_{uuid4().hex[:12]}",
        correlation_id=f"cr_persona_ver_{uuid4().hex[:12]}",
        idempotency_key=f"idem_persona_ver_{persona_pack_id}_{body.version_name}_{uuid4().hex[:8]}",
        persona_pack_id=persona_pack_id,
        version_name=body.version_name,
        style_json=body.style_json,
        voice_json=body.voice_json,
        camera_json=body.camera_json,
    )
    db.add(row)
    db.commit()
    db.refresh(row)
    return PersonaVersionResponse(
        id=row.id,
        persona_pack_id=row.persona_pack_id,
        version_name=row.version_name,
        style_json=row.style_json or {},
        voice_json=row.voice_json or {},
        camera_json=row.camera_json or {},
    )


@router.get("/persona-packs/{persona_pack_id}/versions", response_model=list[PersonaVersionResponse])
def list_persona_versions(persona_pack_id: str, db: Session = Depends(get_db)) -> list[PersonaVersionResponse]:
    rows = db.execute(
        select(PersonaPackVersion)
        .where(
            PersonaPackVersion.persona_pack_id == persona_pack_id,
            PersonaPackVersion.deleted_at.is_(None),
        )
        .order_by(PersonaPackVersion.created_at.desc())
    ).scalars().all()
    return [
        PersonaVersionResponse(
            id=row.id,
            persona_pack_id=row.persona_pack_id,
            version_name=row.version_name,
            style_json=row.style_json or {},
            voice_json=row.voice_json or {},
            camera_json=row.camera_json or {},
        )
        for row in rows
    ]


@router.post("/persona-versions/{persona_pack_version_id}/bindings", response_model=PersonaBindingResponse)
def bind_persona_resources(
    persona_pack_version_id: str,
    body: PersonaBindingRequest,
    db: Session = Depends(get_db),
) -> PersonaBindingResponse:
    version = db.get(PersonaPackVersion, persona_pack_version_id)
    if version is None:
        raise HTTPException(status_code=404, detail="persona version not found")

    dataset_ids: list[str] = []
    for collection_id in body.dataset_collection_ids:
        row = db.execute(
            select(PersonaDatasetBinding).where(
                PersonaDatasetBinding.tenant_id == body.tenant_id,
                PersonaDatasetBinding.project_id == body.project_id,
                PersonaDatasetBinding.persona_pack_version_id == persona_pack_version_id,
                PersonaDatasetBinding.collection_id == collection_id,
                PersonaDatasetBinding.deleted_at.is_(None),
            )
        ).scalars().first()
        if row is None:
            row = PersonaDatasetBinding(
                id=f"pdb_{uuid4().hex}",
                tenant_id=body.tenant_id,
                project_id=body.project_id,
                trace_id=f"tr_pdb_{uuid4().hex[:12]}",
                correlation_id=f"cr_pdb_{uuid4().hex[:12]}",
                idempotency_key=f"idem_pdb_{persona_pack_version_id}_{collection_id}",
                persona_pack_version_id=persona_pack_version_id,
                collection_id=collection_id,
                binding_role=body.binding_role,
                weight=1.0,
            )
            db.add(row)
        dataset_ids.append(collection_id)

    index_ids: list[str] = []
    for kb_version_id in body.kb_version_ids:
        row = db.execute(
            select(PersonaIndexBinding).where(
                PersonaIndexBinding.tenant_id == body.tenant_id,
                PersonaIndexBinding.project_id == body.project_id,
                PersonaIndexBinding.persona_pack_version_id == persona_pack_version_id,
                PersonaIndexBinding.kb_version_id == kb_version_id,
                PersonaIndexBinding.deleted_at.is_(None),
            )
        ).scalars().first()
        if row is None:
            row = PersonaIndexBinding(
                id=f"pib_{uuid4().hex}",
                tenant_id=body.tenant_id,
                project_id=body.project_id,
                trace_id=f"tr_pib_{uuid4().hex[:12]}",
                correlation_id=f"cr_pib_{uuid4().hex[:12]}",
                idempotency_key=f"idem_pib_{persona_pack_version_id}_{kb_version_id}",
                persona_pack_version_id=persona_pack_version_id,
                kb_version_id=kb_version_id,
                priority=100,
                retrieval_policy_json={"binding_role": body.binding_role},
            )
            db.add(row)
        index_ids.append(kb_version_id)

    db.commit()
    return PersonaBindingResponse(
        persona_pack_version_id=persona_pack_version_id,
        dataset_bindings=dataset_ids,
        index_bindings=index_ids,
    )


@router.post("/persona-preview", response_model=PersonaPreviewResponse)
def persona_preview(body: PersonaPreviewRequest, db: Session = Depends(get_db)) -> PersonaPreviewResponse:
    version = db.get(PersonaPackVersion, body.persona_pack_version_id)
    if version is None:
        raise HTTPException(status_code=404, detail="persona version not found")

    dataset_bindings = db.execute(
        select(PersonaDatasetBinding).where(
            PersonaDatasetBinding.tenant_id == body.tenant_id,
            PersonaDatasetBinding.project_id == body.project_id,
            PersonaDatasetBinding.persona_pack_version_id == body.persona_pack_version_id,
            PersonaDatasetBinding.deleted_at.is_(None),
        )
    ).scalars().all()
    index_bindings = db.execute(
        select(PersonaIndexBinding).where(
            PersonaIndexBinding.tenant_id == body.tenant_id,
            PersonaIndexBinding.project_id == body.project_id,
            PersonaIndexBinding.persona_pack_version_id == body.persona_pack_version_id,
            PersonaIndexBinding.deleted_at.is_(None),
        )
    ).scalars().all()

    collection_ids = [item.collection_id for item in dataset_bindings]
    kb_version_ids = [item.kb_version_id for item in index_bindings]

    docs_stmt = select(RagDocument).where(
        RagDocument.tenant_id == body.tenant_id,
        RagDocument.project_id == body.project_id,
        RagDocument.deleted_at.is_(None),
    )
    if collection_ids:
        docs_stmt = docs_stmt.where(RagDocument.collection_id.in_(collection_ids))
    if kb_version_ids:
        docs_stmt = docs_stmt.where(RagDocument.kb_version_id.in_(kb_version_ids))
    docs = db.execute(docs_stmt).scalars().all()

    tokens = [tok for tok in re.split(r"\W+", body.query.lower()) if tok]

    scored: list[tuple[float, RagDocument]] = []
    for doc in docs:
        text = (doc.content_text or "").lower()
        if not tokens:
            score = 0.0
        else:
            score = float(sum(text.count(token) for token in tokens))
        if score <= 0:
            continue
        scored.append((score, doc))
    scored.sort(key=lambda item: item[0], reverse=True)

    chunks: list[PersonaPreviewChunk] = []
    for score, doc in scored[: body.top_k]:
        snippet = (doc.content_text or "")[:240]
        chunks.append(
            PersonaPreviewChunk(
                doc_id=doc.id,
                collection_id=doc.collection_id,
                kb_version_id=doc.kb_version_id,
                title=doc.title,
                source_type=doc.source_type.value if hasattr(doc.source_type, "value") else str(doc.source_type),
                source_id=doc.source_id,
                score=score,
                snippet=snippet,
            )
        )

    return PersonaPreviewResponse(
        persona_pack_version_id=body.persona_pack_version_id,
        query=body.query,
        top_k=body.top_k,
        chunks=chunks,
    )


@router.post("/knowledge-packs/bootstrap", response_model=KnowledgePackBootstrapResponse, status_code=201)
def bootstrap_knowledge_pack(
    body: KnowledgePackBootstrapRequest,
    db: Session = Depends(get_db),
) -> KnowledgePackBootstrapResponse:
    role_id = body.role_id.strip().lower()
    if not role_id:
        raise HTTPException(status_code=400, detail="REQ-VALIDATION-001: role_id required")

    pack_name = (body.pack_name or f"{role_id}_bootstrap_pack").strip()
    if not pack_name:
        raise HTTPException(status_code=400, detail="REQ-VALIDATION-001: pack_name required")

    collection = db.execute(
        select(RagCollection).where(
            RagCollection.tenant_id == body.tenant_id,
            RagCollection.project_id == body.project_id,
            RagCollection.name == pack_name,
            RagCollection.language_code == body.language_code,
            RagCollection.deleted_at.is_(None),
        )
    ).scalars().first()
    if collection is None:
        collection = RagCollection(
            id=f"rag_collection_{uuid4().hex}",
            tenant_id=body.tenant_id,
            project_id=body.project_id,
            trace_id=f"tr_bootstrap_collection_{uuid4().hex[:12]}",
            correlation_id=f"cr_bootstrap_collection_{uuid4().hex[:12]}",
            idempotency_key=f"idem_bootstrap_collection_{pack_name}_{uuid4().hex[:8]}",
            novel_id=None,
            name=pack_name,
            version="v1",
            language_code=body.language_code,
            description=f"bootstrap knowledge pack for {role_id}",
            tags_json=["bootstrap", role_id],
        )
        db.add(collection)
        db.flush()

    kb_version = KbVersion(
        id=f"kb_version_{uuid4().hex}",
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        trace_id=f"tr_bootstrap_kb_{uuid4().hex[:12]}",
        correlation_id=f"cr_bootstrap_kb_{uuid4().hex[:12]}",
        idempotency_key=f"idem_bootstrap_kb_{collection.id}_{uuid4().hex[:8]}",
        collection_id=collection.id,
        version_name=f"bootstrap_{datetime.now(timezone.utc).strftime('%Y%m%d%H%M%S')}",
        status="released",
        recipe_id=f"bootstrap_{role_id}",
        embedding_model_profile_id=None,
        release_note=f"bootstrap knowledge template for {role_id}",
    )
    db.add(kb_version)
    db.flush()

    sections = _knowledge_template_sections(role_id, body.template_key)
    scope = _scope_from_text(body.default_knowledge_scope)
    created_documents = 0
    chunk_count = 0
    all_text_parts: list[str] = []
    pack_id = f"kp_{uuid4().hex}"
    for section in sections:
        section_title = str(section.get("title") or "knowledge")
        section_source_type = section.get("source_type") or RagSourceType.rule
        section_content = str(section.get("content") or "")
        chunks = _chunk_text(section_content)
        chunk_count += len(chunks)
        for idx, chunk in enumerate(chunks):
            fingerprint = hashlib.sha1(chunk.encode("utf-8")).hexdigest()
            db.add(
                RagDocument(
                    id=f"rag_doc_{uuid4().hex}",
                    tenant_id=body.tenant_id,
                    project_id=body.project_id,
                    trace_id=f"tr_bootstrap_doc_{uuid4().hex[:12]}",
                    correlation_id=f"cr_bootstrap_doc_{uuid4().hex[:12]}",
                    idempotency_key=f"idem_bootstrap_doc_{pack_id}_{section_title}_{idx}_{fingerprint[:8]}",
                    collection_id=collection.id,
                    kb_version_id=kb_version.id,
                    novel_id=None,
                    scope=scope,
                    source_type=section_source_type,
                    source_id=pack_id,
                    language_code=body.language_code,
                    title=section_title,
                    content_text=chunk,
                    metadata_json={
                        "pack_id": pack_id,
                        "role_id": role_id,
                        "template_key": body.template_key or "",
                        "section_title": section_title,
                        "chunk_index": idx,
                        "fingerprint": fingerprint,
                    },
                )
            )
            created_documents += 1
            all_text_parts.append(chunk)

    extracted_terms = _extract_terms("\n".join(all_text_parts))
    updated_at = datetime.now(timezone.utc).isoformat()
    stack_payload = {
        "type": "knowledge_pack",
        "pack_id": pack_id,
        "role_id": role_id,
        "pack_name": pack_name,
        "collection_id": collection.id,
        "kb_version_id": kb_version.id,
        "scope": scope.value,
        "status": "ready",
        "created_documents": created_documents,
        "chunk_count": chunk_count,
        "extracted_terms": extracted_terms,
        "schema_version": "1.0",
        "updated_at": updated_at,
    }
    _upsert_stack(
        db,
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        stack_name=_knowledge_pack_stack_name(pack_id),
        payload=stack_payload,
        trace_prefix="knowledge_pack",
    )
    db.commit()

    return KnowledgePackBootstrapResponse(
        pack_id=pack_id,
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        role_id=role_id,
        pack_name=pack_name,
        collection_id=collection.id,
        kb_version_id=kb_version.id,
        scope=scope.value,
        created_documents=created_documents,
        chunk_count=chunk_count,
        extracted_terms=extracted_terms,
        updated_at=updated_at,
    )


@router.get("/knowledge-packs", response_model=list[KnowledgePackItemResponse])
def list_knowledge_packs(
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    role_id: str | None = Query(default=None),
    db: Session = Depends(get_db),
) -> list[KnowledgePackItemResponse]:
    rows = db.execute(
        select(CreativePolicyStack)
        .where(
            CreativePolicyStack.tenant_id == tenant_id,
            CreativePolicyStack.project_id == project_id,
            CreativePolicyStack.name.like("knowledge_pack:%"),
            CreativePolicyStack.deleted_at.is_(None),
        )
        .order_by(CreativePolicyStack.created_at.desc())
    ).scalars().all()
    result: list[KnowledgePackItemResponse] = []
    for row in rows:
        pack_id = row.name.split(":", 1)[1]
        payload = dict(row.stack_json or {})
        if role_id and str(payload.get("role_id") or "").strip().lower() != role_id.strip().lower():
            continue
        result.append(
            _build_knowledge_pack_item(
                payload,
                pack_id=pack_id,
                tenant_id=tenant_id,
                project_id=project_id,
            )
        )
    return result


@router.post("/import-jobs", response_model=KnowledgeImportJobResponse, status_code=201)
def create_import_job(
    body: KnowledgeImportRequest,
    db: Session = Depends(get_db),
) -> KnowledgeImportJobResponse:
    source_name = body.source_name.strip()
    if not source_name:
        raise HTTPException(status_code=400, detail="REQ-VALIDATION-001: source_name required")
    if not body.content_text.strip():
        raise HTTPException(status_code=400, detail="REQ-VALIDATION-001: content_text required")

    collection = db.execute(
        select(RagCollection).where(
            RagCollection.id == body.collection_id,
            RagCollection.tenant_id == body.tenant_id,
            RagCollection.project_id == body.project_id,
            RagCollection.deleted_at.is_(None),
        )
    ).scalars().first()
    if collection is None:
        raise HTTPException(status_code=404, detail="REQ-VALIDATION-001: collection not found")

    kb_version_id = body.kb_version_id
    kb_version = None
    if kb_version_id:
        kb_version = db.execute(
            select(KbVersion).where(
                KbVersion.id == kb_version_id,
                KbVersion.collection_id == body.collection_id,
                KbVersion.deleted_at.is_(None),
            )
        ).scalars().first()
        if kb_version is None:
            raise HTTPException(status_code=404, detail="REQ-VALIDATION-001: kb_version not found")
    else:
        kb_version = db.execute(
            select(KbVersion)
            .where(
                KbVersion.collection_id == body.collection_id,
                KbVersion.deleted_at.is_(None),
            )
            .order_by(KbVersion.created_at.desc())
        ).scalars().first()
        if kb_version is not None:
            kb_version_id = kb_version.id

    normalized_text = _normalize_import_content(body.content_text, body.source_format)
    chunks = _chunk_text(normalized_text)
    scope = _scope_from_text(body.scope)
    source_type = _source_type_from_format(body.source_format)
    import_job_id = f"import_{uuid4().hex}"

    existing_docs = db.execute(
        select(RagDocument).where(
            RagDocument.tenant_id == body.tenant_id,
            RagDocument.project_id == body.project_id,
            RagDocument.collection_id == body.collection_id,
            RagDocument.deleted_at.is_(None),
        )
    ).scalars().all()
    existing_fingerprints = {
        hashlib.sha1((doc.content_text or "").strip().encode("utf-8")).hexdigest()
        for doc in existing_docs
    }

    created_documents = 0
    deduplicated_documents = 0
    for idx, chunk in enumerate(chunks):
        fingerprint = hashlib.sha1(chunk.encode("utf-8")).hexdigest()
        if fingerprint in existing_fingerprints:
            deduplicated_documents += 1
            continue
        existing_fingerprints.add(fingerprint)
        db.add(
            RagDocument(
                id=f"rag_doc_{uuid4().hex}",
                tenant_id=body.tenant_id,
                project_id=body.project_id,
                trace_id=f"tr_rag_import_{uuid4().hex[:12]}",
                correlation_id=f"cr_rag_import_{uuid4().hex[:12]}",
                idempotency_key=f"idem_rag_import_{import_job_id}_{idx}_{fingerprint[:8]}",
                collection_id=body.collection_id,
                kb_version_id=kb_version_id,
                novel_id=collection.novel_id,
                scope=scope,
                source_type=source_type,
                source_id=f"{source_name}:{idx}",
                language_code=body.language_code,
                title=source_name,
                content_text=chunk,
                metadata_json={
                    "import_job_id": import_job_id,
                    "source_format": body.source_format.strip().lower(),
                    "chunk_index": idx,
                    "fingerprint": fingerprint,
                    "role_ids": body.role_ids,
                },
            )
        )
        created_documents += 1

    extracted_terms = _extract_terms(normalized_text)
    updated_at = datetime.now(timezone.utc).isoformat()
    knowledge_change_report = {
        "added_documents": created_documents,
        "deduplicated_documents": deduplicated_documents,
        "total_chunks": len(chunks),
        "affected_roles": body.role_ids,
        "source_name": source_name,
        "source_format": body.source_format.strip().lower(),
        "new_terms": extracted_terms,
    }
    stack_payload = {
        "type": "rag_import_job",
        "import_job_id": import_job_id,
        "collection_id": body.collection_id,
        "kb_version_id": kb_version_id,
        "source_name": source_name,
        "source_format": body.source_format.strip().lower(),
        "status": "completed",
        "created_documents": created_documents,
        "deduplicated_documents": deduplicated_documents,
        "chunk_count": len(chunks),
        "extracted_terms": extracted_terms,
        "affected_roles": body.role_ids,
        "knowledge_change_report": knowledge_change_report,
        "schema_version": "1.0",
        "updated_at": updated_at,
    }
    _upsert_stack(
        db,
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        stack_name=_import_job_stack_name(import_job_id),
        payload=stack_payload,
        trace_prefix="rag_import_job",
    )
    db.commit()
    response = _build_import_job_item(
        stack_payload,
        import_job_id=import_job_id,
        tenant_id=body.tenant_id,
        project_id=body.project_id,
    )
    notify_telegram_event(
        db=db,
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        event_type="rag.embedding.completed",
        summary="RAG import completed",
        extra={
            "import_job_id": import_job_id,
            "collection_id": body.collection_id,
            "kb_version_id": kb_version_id,
            "source_name": source_name,
            "source_format": body.source_format.strip().lower(),
            "created_documents": response.created_documents,
            "deduplicated_documents": response.deduplicated_documents,
            "chunk_count": response.chunk_count,
        },
    )
    return response


@router.get("/import-jobs", response_model=list[KnowledgeImportJobResponse])
def list_import_jobs(
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    role_id: str | None = Query(default=None),
    db: Session = Depends(get_db),
) -> list[KnowledgeImportJobResponse]:
    rows = db.execute(
        select(CreativePolicyStack)
        .where(
            CreativePolicyStack.tenant_id == tenant_id,
            CreativePolicyStack.project_id == project_id,
            CreativePolicyStack.name.like("rag_import_job:%"),
            CreativePolicyStack.deleted_at.is_(None),
        )
        .order_by(CreativePolicyStack.created_at.desc())
    ).scalars().all()
    result: list[KnowledgeImportJobResponse] = []
    role_filter = role_id.strip().lower() if role_id else None
    for row in rows:
        import_job_id = row.name.split(":", 1)[1]
        payload = dict(row.stack_json or {})
        if role_filter:
            affected_roles = [str(item).strip().lower() for item in list(payload.get("affected_roles") or [])]
            if role_filter not in affected_roles:
                continue
        result.append(
            _build_import_job_item(
                payload,
                import_job_id=import_job_id,
                tenant_id=tenant_id,
                project_id=project_id,
            )
        )
    return result


@router.delete("/collections/{collection_id}")
def delete_collection(
    collection_id: str,
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    db: Session = Depends(get_db),
) -> dict[str, str]:
    row = db.execute(
        select(RagCollection).where(
            RagCollection.id == collection_id,
            RagCollection.tenant_id == tenant_id,
            RagCollection.project_id == project_id,
            RagCollection.deleted_at.is_(None),
        )
    ).scalars().first()
    if row is None:
        raise HTTPException(status_code=404, detail="REQ-VALIDATION-001: collection not found")
    row.deleted_at = datetime.now(timezone.utc)
    db.commit()
    return {"status": "deleted", "id": collection_id}


@router.delete("/persona-packs/{persona_pack_id}")
def delete_persona_pack(
    persona_pack_id: str,
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    db: Session = Depends(get_db),
) -> dict[str, str]:
    row = db.execute(
        select(PersonaPack).where(
            PersonaPack.id == persona_pack_id,
            PersonaPack.tenant_id == tenant_id,
            PersonaPack.project_id == project_id,
            PersonaPack.deleted_at.is_(None),
        )
    ).scalars().first()
    if row is None:
        raise HTTPException(status_code=404, detail="REQ-VALIDATION-001: persona pack not found")
    row.deleted_at = datetime.now(timezone.utc)
    db.commit()
    return {"status": "deleted", "id": persona_pack_id}


# TASK_CARD_26_BINARY_IMPORT_ASYNC: 二进制文件导入
_binary_import_jobs: dict[str, dict] = {}  # 简单内存存储


def _extract_text_from_binary(file_name: str, file_bytes: bytes) -> tuple[str, int, int, int]:
    """
    从二进制文件抽取文本（支持 pdf/xlsx/docx/txt）

    返回: (text, pages, tables, images)
    """
    import io
    file_lower = file_name.lower()

    try:
        if file_lower.endswith('.txt'):
            # Try common encodings
            for enc in ("utf-8", "gbk", "gb2312", "latin-1"):
                try:
                    return file_bytes.decode(enc), 1, 0, 0
                except UnicodeDecodeError:
                    continue
            return file_bytes.decode('utf-8', errors='ignore'), 1, 0, 0

        elif file_lower.endswith('.pdf'):
            try:
                import pypdf  # type: ignore[import]
                reader = pypdf.PdfReader(io.BytesIO(file_bytes))
                pages_count = len(reader.pages)
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
                return "\n\n".join(text_parts), pages_count, 0, 0
            except ImportError:
                # Fallback: minimal text extraction from raw PDF bytes
                import re as _re
                raw = file_bytes.decode('latin-1', errors='ignore')
                text_blocks = _re.findall(r'BT\s*(.*?)\s*ET', raw, _re.DOTALL)
                extracted = " ".join(text_blocks)[:8000]
                return extracted or "PDF parsed (pypdf not installed)", 1, 0, 0
            except Exception as e:
                return f"PDF parsing error: {e}", 1, 0, 0

        elif file_lower.endswith(('.xlsx', '.xls')):
            try:
                import openpyxl  # type: ignore[import]
                wb = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
                text_parts = []
                table_count = 0
                for sheet in wb.sheetnames:
                    ws = wb[sheet]
                    rows_text = []
                    for row in ws.iter_rows(values_only=True):
                        cells = [str(cell) if cell is not None else "" for cell in row]
                        row_text = "\t".join(cells).strip()
                        if row_text:
                            rows_text.append(row_text)
                    if rows_text:
                        text_parts.append(f"[Sheet: {sheet}]\n" + "\n".join(rows_text[:200]))
                        table_count += 1
                return "\n\n".join(text_parts), 1, table_count, 0
            except ImportError:
                return f"Excel file received ({len(file_bytes)} bytes). Install openpyxl for full parsing.", 1, 1, 0
            except Exception as e:
                return f"Excel parsing error: {e}", 1, 0, 0

        elif file_lower.endswith('.docx'):
            try:
                import docx  # type: ignore[import]
                doc = docx.Document(io.BytesIO(file_bytes))
                paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
                table_count = len(doc.tables)
                table_texts = []
                for table in doc.tables:
                    rows = []
                    for row in table.rows:
                        cells = [cell.text.strip() for cell in row.cells]
                        rows.append("\t".join(cells))
                    table_texts.append("\n".join(rows))
                all_text = "\n\n".join(paragraphs)
                if table_texts:
                    all_text += "\n\n" + "\n\n".join(table_texts)
                return all_text, 1, table_count, 0
            except ImportError:
                return f"DOCX file received ({len(file_bytes)} bytes). Install python-docx for full parsing.", 1, 0, 0
            except Exception as e:
                return f"DOCX parsing error: {e}", 1, 0, 0

        else:
            return "Unsupported file format", 0, 0, 0

    except Exception as e:
        return f"Error extracting text: {e}", 0, 0, 0


@router.post("/collections/{collection_id}/binary-import", response_model=BinaryImportJobResponse, status_code=202)
async def upload_binary_file_for_import(
    collection_id: str,
    file: UploadFile = File(...),
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    model_provider_id: str | None = Query(default=None),
    use_vision: bool = Query(default=False),
    db: Session = Depends(get_db),
) -> BinaryImportJobResponse:
    """上传二进制文件（PDF/XLSX/DOCX/TXT）进行异步导入，支持多模态LLM提取"""
    collection = db.execute(
        select(RagCollection).where(
            RagCollection.id == collection_id,
            RagCollection.tenant_id == tenant_id,
            RagCollection.project_id == project_id,
            RagCollection.deleted_at.is_(None),
        )
    ).scalars().first()
    
    if not collection:
        raise HTTPException(status_code=404, detail="collection not found")
    
    # 检查文件格式
    allowed_formats = {'.pdf', '.xlsx', '.xls', '.docx', '.txt'}
    file_ext = ''.join(f for f in ['.'.join(file.filename.split('.')[1:])] if f)
    if not any(file.filename.lower().endswith(fmt) for fmt in allowed_formats):
        raise HTTPException(status_code=400, detail="unsupported file format")
    
    # 读取文件内容
    file_bytes = await file.read()
    file_size = len(file_bytes)
    
    if file_size > 50_000_000:  # 50MB limit
        raise HTTPException(status_code=413, detail="file too large (max 50MB)")
    
    import_job_id = f"import_bin_{uuid4().hex[:12]}"
    
    # 提取文本
    extracted_text, pages, tables, images = _extract_text_from_binary(file.filename, file_bytes)

    # 多模态 LLM 增强提取（可选）
    if use_vision and model_provider_id:
        try:
            provider = db.get(ModelProvider, model_provider_id)
            if provider and provider.deleted_at is None:
                provider_settings = _load_provider_settings_rag(
                    db,
                    tenant_id=tenant_id,
                    project_id=project_id,
                    provider_id=model_provider_id,
                )
                vision_prompt = (
                    f"请从以下文档内容中提取和整理关键信息，输出为结构化文本：\n\n{extracted_text[:4000]}"
                )
                enhanced = _call_provider_llm(
                    provider=provider,
                    provider_settings=provider_settings,
                    prompt=vision_prompt,
                    max_tokens=1500,
                )
                if enhanced:
                    extracted_text = enhanced
        except (requests.RequestException, ValueError):
            pass  # 失败时使用原始提取文本

    # 存储到内存
    _binary_import_jobs[import_job_id] = {
        "tenant_id": tenant_id,
        "project_id": project_id,
        "collection_id": collection_id,
        "file_name": file.filename,
        "file_format": file_ext.lower(),
        "file_size_bytes": file_size,
        "status": "processing",
        "progress_percent": 50,
        "extracted_text": extracted_text,
        "extracted_pages": pages,
        "extracted_tables": tables,
        "extracted_images": images,
        "created_at": datetime.now(timezone.utc),
        "updated_at": datetime.now(timezone.utc),
    }
    
    # 记录事件
    event = WorkflowEvent(
        id=f"evt_{uuid4().hex[:24]}",
        tenant_id=tenant_id,
        project_id=project_id,
        trace_id=f"tr_bin_import_{uuid4().hex[:12]}",
        correlation_id=f"cr_bin_import_{uuid4().hex[:12]}",
        idempotency_key=f"idem_bin_import_{import_job_id}",
        run_id=None,
        stage=None,
        event_type="rag.binary.import.started",
        event_version="1.0",
        producer="studio_api",
        occurred_at=datetime.now(timezone.utc),
        payload_json={
            "import_job_id": import_job_id,
            "file_name": file.filename,
            "file_size": file_size,
            "collection_id": collection_id,
        },
    )
    try:
        from ainern2d_shared.ainer_db_models.pipeline_models import WorkflowEvent
        db.add(event)
        db.commit()
    except Exception:
        pass  # 事件记录失败不影响导入
    
    return BinaryImportJobResponse(
        import_job_id=import_job_id,
        tenant_id=tenant_id,
        project_id=project_id,
        collection_id=collection_id,
        file_name=file.filename,
        file_format=file_ext.lower(),
        file_size_bytes=file_size,
        status="processing",
        progress_percent=50,
        extracted_text_preview=extracted_text[:200],
        extracted_pages=pages,
        extracted_tables=tables,
        extracted_images=images,
    )


@router.get("/binary-import/{import_job_id}", response_model=BinaryImportJobResponse)
def get_binary_import_status(import_job_id: str) -> BinaryImportJobResponse:
    """查询二进制导入任务状态"""
    job = _binary_import_jobs.get(import_job_id)
    if not job:
        raise HTTPException(status_code=404, detail="import job not found")
    
    return BinaryImportJobResponse(
        import_job_id=import_job_id,
        tenant_id=job["tenant_id"],
        project_id=job["project_id"],
        collection_id=job["collection_id"],
        file_name=job["file_name"],
        file_format=job["file_format"],
        file_size_bytes=job["file_size_bytes"],
        status=job["status"],
        progress_percent=job["progress_percent"],
        extracted_text_preview=job["extracted_text"][:200],
        extracted_pages=job["extracted_pages"],
        extracted_tables=job["extracted_tables"],
        extracted_images=job["extracted_images"],
    )


class NovelRagInitRequest(BaseModel):
    tenant_id: str
    project_id: str
    model_provider_id: str
    novel_id: str
    max_tokens: int = Field(default=1200, ge=200, le=3000)


class NovelRagInitResponse(BaseModel):
    collection_id: str
    novel_id: str
    documents_created: int
    chunks_total: int
    status: str


def _load_provider_settings_rag(
    db: Session,
    *,
    tenant_id: str,
    project_id: str,
    provider_id: str,
) -> dict:
    row = db.execute(
        select(CreativePolicyStack).where(
            CreativePolicyStack.tenant_id == tenant_id,
            CreativePolicyStack.project_id == project_id,
            CreativePolicyStack.name == f"provider_settings:{provider_id}",
            CreativePolicyStack.deleted_at.is_(None),
        )
    ).scalars().first()
    if row is None:
        return {}
    return dict(row.stack_json or {})


def _call_provider_llm(
    *,
    provider: ModelProvider,
    provider_settings: dict,
    prompt: str,
    max_tokens: int,
) -> str:
    endpoint = (provider.endpoint or "").strip().rstrip("/")
    token = str(provider_settings.get("access_token") or "").strip()
    model_catalog = list(provider_settings.get("model_catalog") or [])
    model_name = model_catalog[0] if model_catalog else "gpt-4o-mini"
    if not endpoint or not token:
        raise ValueError("missing provider endpoint/token")

    payload = {
        "model": model_name,
        "messages": [
            {"role": "system", "content": "你是专业知识提取助手。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.3,
        "max_tokens": max_tokens,
    }
    response = requests.post(
        f"{endpoint}/chat/completions",
        json=payload,
        headers={
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json",
        },
        timeout=90.0,
    )
    if response.status_code < 200 or response.status_code >= 300:
        raise ValueError(f"provider_http_status_{response.status_code}")
    parsed = response.json() if response.text else {}
    choices = parsed.get("choices") or []
    if not choices:
        raise ValueError("provider_response_missing_choices")
    return str((choices[0].get("message") or {}).get("content") or "").strip()


_NOVEL_RAG_PROMPT = """请从以下小说章节内容中提取关键知识点，以便用于AI辅助创作：
- 主要人物的性格、背景、关系
- 重要地点和场景的描述
- 关键情节节点和伏笔
- 世界观规则和设定

请将内容整理为结构化段落，每个知识点单独一段。

章节内容：
{content}
"""


@router.post("/collections/{collection_id}/novel-init", response_model=NovelRagInitResponse)
def init_novel_rag(
    collection_id: str,
    body: NovelRagInitRequest,
    db: Session = Depends(get_db),
) -> NovelRagInitResponse:
    collection = db.get(RagCollection, collection_id)
    if collection is None:
        raise HTTPException(status_code=404, detail="RAG collection not found")

    novel = db.get(Novel, body.novel_id)
    if novel is None:
        raise HTTPException(status_code=404, detail="novel not found")

    provider = db.get(ModelProvider, body.model_provider_id)
    if provider is None or provider.deleted_at is not None:
        raise HTTPException(status_code=404, detail="model provider not found")

    chapters = db.execute(
        select(Chapter).where(
            Chapter.novel_id == body.novel_id,
            Chapter.deleted_at.is_(None),
        ).order_by(Chapter.chapter_no.asc())
    ).scalars().all()

    if not chapters:
        raise HTTPException(status_code=404, detail="no chapters found in novel")

    provider_settings = _load_provider_settings_rag(
        db,
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        provider_id=provider.id,
    )

    documents_created = 0
    chunks_total = 0
    now = datetime.now(timezone.utc)

    for ch in chapters:
        raw_text = (ch.raw_text or "").strip()
        if not raw_text:
            continue

        # Try LLM extraction; fall back to direct chunking
        try:
            prompt = _NOVEL_RAG_PROMPT.format(content=raw_text[:4000])
            extracted_text = _call_provider_llm(
                provider=provider,
                provider_settings=provider_settings,
                prompt=prompt,
                max_tokens=body.max_tokens,
            )
        except (requests.RequestException, ValueError):
            extracted_text = raw_text

        chunks = _chunk_text(extracted_text, chunk_size=480)

        for idx, chunk in enumerate(chunks):
            if not chunk.strip():
                continue
            fingerprint = hashlib.sha1(chunk.encode("utf-8")).hexdigest()
            doc = RagDocument(
                id=f"doc_novel_{uuid4().hex}",
                tenant_id=body.tenant_id,
                project_id=body.project_id,
                trace_id=f"tr_novel_rag_{uuid4().hex[:12]}",
                correlation_id=f"cr_novel_rag_{uuid4().hex[:12]}",
                idempotency_key=f"idem_novel_rag_{ch.id}_{idx}_{fingerprint[:8]}",
                collection_id=collection_id,
                novel_id=body.novel_id,
                scope=RagScope.novel,
                source_type=RagSourceType.chapter,
                source_id=ch.id,
                title=f"{novel.title} - 第{ch.chapter_no}章 [{idx + 1}]",
                content_text=chunk,
                language_code=ch.language_code or "zh",
                metadata_json={"chapter_no": ch.chapter_no, "chunk_index": idx},
            )
            db.add(doc)
            chunks_total += 1

        documents_created += 1

    db.commit()

    notify_telegram_event(
        db=db,
        tenant_id=body.tenant_id,
        project_id=body.project_id,
        event_type="plan.prompt.generated",
        summary="Novel RAG initialization completed",
        trace_id=novel.trace_id,
        correlation_id=novel.correlation_id,
        extra={
            "collection_id": collection_id,
            "novel_id": body.novel_id,
            "documents_created": documents_created,
            "chunks_total": chunks_total,
        },
    )

    return NovelRagInitResponse(
        collection_id=collection_id,
        novel_id=body.novel_id,
        documents_created=documents_created,
        chunks_total=chunks_total,
        status="completed",
    )


# ═══════════════════════════════════════════════════════════════════════════════
# Effective KB — 三层 KB 合并接口（novel > role > persona）
# ═══════════════════════════════════════════════════════════════════════════════

class EffectiveKBItem(BaseModel):
    collection_id: str
    collection_name: str
    bind_type: str
    bind_id: str | None = None
    priority: int  # 1=novel, 2=role, 3=persona, 4=global


class EffectiveKBResponse(BaseModel):
    persona_pack_id: str
    role_id: str | None = None
    novel_id: str | None = None
    effective_collections: list[EffectiveKBItem] = Field(default_factory=list)


@router.get("/effective-kb", response_model=EffectiveKBResponse)
def get_effective_kb(
    tenant_id: str = Query(...),
    project_id: str = Query(...),
    persona_pack_id: str = Query(...),
    novel_id: str | None = Query(default=None),
    db: Session = Depends(get_db),
) -> EffectiveKBResponse:
    """计算 Persona 的生效 KB 列表（novel > role > persona > global）"""
    persona = db.get(PersonaPack, persona_pack_id)
    if persona is None:
        raise HTTPException(status_code=404, detail="persona pack not found")

    role_id = persona.role_id
    seen_names: set[str] = set()
    items: list[EffectiveKBItem] = []

    def _collect(bt: KBBindType, bid: str | None, priority: int) -> None:
        if bid is None and bt != KBBindType.global_:
            return
        stmt = select(RagCollection).where(
            RagCollection.tenant_id == tenant_id,
            RagCollection.project_id == project_id,
            RagCollection.bind_type == bt,
            RagCollection.deleted_at.is_(None),
        )
        if bt != KBBindType.global_:
            stmt = stmt.where(RagCollection.bind_id == bid)
        rows = db.execute(stmt.order_by(RagCollection.created_at.desc())).scalars().all()
        for row in rows:
            if row.name not in seen_names:
                seen_names.add(row.name)
                items.append(EffectiveKBItem(
                    collection_id=row.id,
                    collection_name=row.name,
                    bind_type=row.bind_type.value if row.bind_type else "unknown",
                    bind_id=row.bind_id,
                    priority=priority,
                ))

    # 优先级 1: Novel KB（项目一致性第一）
    if novel_id:
        _collect(KBBindType.novel, novel_id, 1)
    # 优先级 2: Role KB（职业专业知识）
    if role_id:
        _collect(KBBindType.role, role_id, 2)
    # 优先级 3: Persona KB（个人风格）
    _collect(KBBindType.persona, persona_pack_id, 3)
    # 优先级 4: Global KB（全局知识）
    _collect(KBBindType.global_, None, 4)

    items.sort(key=lambda x: x.priority)

    return EffectiveKBResponse(
        persona_pack_id=persona_pack_id,
        role_id=role_id,
        novel_id=novel_id,
        effective_collections=items,
    )
